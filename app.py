import streamlit as st
import plotly.graph_objects as go
from io import BytesIO
from docx import Document
from docx.shared import Pt, RGBColor, Inches
import json
from datetime import datetime

# --- Page Configuration ---
st.set_page_config(
    page_title="Fintech IMPACT Radar",
    page_icon="📡",
    layout="wide",
    initial_sidebar_state="expanded"
)

# --- Enhanced Custom CSS ---
st.markdown("""
    <style>
    /* Import Google Font */
    @import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;500;600;700&display=swap');
    
    /* Main Background with gradient */
    .stApp {
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        font-family: 'Inter', sans-serif;
    }
    
    /* Main content area */
    .main .block-container {
        padding-top: 2rem;
        padding-bottom: 2rem;
        max-width: 1400px;
    }
    
    /* Sidebar styling */
    [data-testid="stSidebar"] {
        background: linear-gradient(180deg, #1e293b 0%, #334155 100%);
    }
    
    [data-testid="stSidebar"] h1, 
    [data-testid="stSidebar"] h2, 
    [data-testid="stSidebar"] h3,
    [data-testid="stSidebar"] label,
    [data-testid="stSidebar"] p {
        color: #f1f5f9 !important;
    }
    
    [data-testid="stSidebar"] .stMarkdown {
        color: #cbd5e1;
    }
    
    /* Card Styling with glassmorphism */
    div[data-testid="stVerticalBlock"] > div[style*="flex-direction: column;"] > div[data-testid="stVerticalBlock"] {
        background: rgba(255, 255, 255, 0.95);
        backdrop-filter: blur(10px);
        padding: 1.5rem;
        border-radius: 12px;
        box-shadow: 0 8px 32px rgba(0, 0, 0, 0.1);
        border: 1px solid rgba(255, 255, 255, 0.2);
    }
    
    /* Headers with gradient text */
    h1 {
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
        background-clip: text;
        font-weight: 700;
        font-size: 3rem !important;
        margin-bottom: 0.5rem;
    }
    
    h2, h3 { 
        color: #1e293b;
        font-weight: 600;
    }
    
    /* Metric styling */
    div[data-testid="stMetricValue"] {
        font-size: 2rem;
        font-weight: 700;
        color: #667eea;
    }
    
    /* Button styling */
    .stButton > button {
        border-radius: 8px;
        font-weight: 600;
        transition: all 0.3s ease;
        border: none;
    }
    
    .stButton > button:hover {
        transform: translateY(-2px);
        box-shadow: 0 4px 12px rgba(0, 0, 0, 0.15);
    }
    
    /* Download button special styling */
    .stDownloadButton > button {
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        color: white;
    }
    
    /* Container borders */
    [data-testid="stVerticalBlock"] [data-testid="stVerticalBlock"] {
        background: white;
        border-radius: 12px;
        padding: 1.25rem;
        box-shadow: 0 2px 8px rgba(0, 0, 0, 0.08);
        transition: all 0.3s ease;
    }
    
    [data-testid="stVerticalBlock"] [data-testid="stVerticalBlock"]:hover {
        box-shadow: 0 4px 16px rgba(102, 126, 234, 0.2);
        transform: translateY(-2px);
    }
    
    /* Slider styling */
    .stSlider { 
        padding-top: 0.5rem; 
        padding-bottom: 0.5rem; 
    }
    
    /* Expander styling */
    .streamlit-expanderHeader {
        background: linear-gradient(135deg, #f1f5f9 0%, #e2e8f0 100%);
        border-radius: 8px;
        font-size: 0.9rem;
        font-weight: 600;
        color: #334155;
        border: 1px solid #cbd5e1;
    }
    
    .streamlit-expanderHeader:hover {
        background: linear-gradient(135deg, #e2e8f0 0%, #cbd5e1 100%);
    }
    
    /* Rubric Table Styling */
    .rubric-table { 
        font-size: 0.85rem; 
        width: 100%; 
        border-collapse: collapse; 
        font-family: 'Inter', sans-serif;
        margin-top: 0.5rem;
    }
    .rubric-table td { 
        padding: 10px; 
        border-bottom: 1px solid #e2e8f0; 
        vertical-align: top; 
        color: #334155;
    }
    .rubric-header { 
        font-weight: 700; 
        color: #1e293b; 
        background: linear-gradient(135deg, #f8fafc 0%, #f1f5f9 100%);
        border-right: 2px solid #667eea;
        width: 80px;
    }
    
    /* Score badge styling */
    .score-badge {
        display: inline-block;
        padding: 4px 12px;
        border-radius: 20px;
        font-weight: 700;
        font-size: 1.1rem;
    }
    
    /* Input fields */
    .stTextInput > div > div > input {
        border-radius: 8px;
        border: 2px solid #e2e8f0;
        background: rgba(255, 255, 255, 0.9);
    }
    
    .stTextInput > div > div > input:focus {
        border-color: #667eea;
        box-shadow: 0 0 0 3px rgba(102, 126, 234, 0.1);
    }
    
    /* Text area */
    .stTextArea textarea {
        border-radius: 8px;
        border: 2px solid #e2e8f0;
    }
    
    .stTextArea textarea:focus {
        border-color: #667eea;
        box-shadow: 0 0 0 3px rgba(102, 126, 234, 0.1);
    }
    
    /* Caption styling */
    .caption {
        font-size: 0.8rem;
        color: #64748b;
        font-weight: 500;
    }
    
    /* Overall score display */
    .overall-score {
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        color: white;
        padding: 2rem;
        border-radius: 16px;
        text-align: center;
        box-shadow: 0 8px 32px rgba(102, 126, 234, 0.3);
        margin-top: 2rem;
    }
    
    /* Info badges */
    .info-badge {
        display: inline-block;
        background: #e0e7ff;
        color: #4338ca;
        padding: 4px 12px;
        border-radius: 12px;
        font-size: 0.75rem;
        font-weight: 600;
        margin: 2px;
    }
    </style>
""", unsafe_allow_html=True)

# --- Data Definitions ---
DIMENSIONS = [
    {
        'id': 'integration',
        'letter': 'I',
        'icon': '🔗',
        'title': 'INTEGRATION',
        'subtitle': 'Connectivity',
        'question': 'Is it an Island or an Ecosystem?',
        'leftLabel': 'Closed / Island',
        'rightLabel': 'Open API Platform',
        'rubric': {
            'low': 'Closed system. No APIs. Hard to export data. "Walled Garden."',
            'medium': 'Some integrations (e.g., connects to Xero), but largely self-contained.',
            'high': 'API-first architecture. Allows developers to build on top. Two-way data flow.'
        },
        'challengePrompts': ['Evidence of API documentation?', 'Are integrations bidirectional?']
    },
    {
        'id': 'monetization',
        'letter': 'M',
        'icon': '💰',
        'title': 'MONETIZATION',
        'subtitle': 'Unit Economics',
        'question': 'Growth at all costs or sustainable?',
        'leftLabel': 'Burning Cash',
        'rightLabel': 'Sustainable Profit',
        'rubric': {
            'low': 'Freemium with no clear upsell. High burn. Subsidized by VC money.',
            'medium': 'Generating revenue (interchange fees), but barely covering costs.',
            'high': 'Strong LTV > CAC. Diversified revenue (Sub + Trans + Data).'
        },
        'challengePrompts': ['Path to profitability?', 'LTV vs CAC ratio?']
    },
    {
        'id': 'painPoint',
        'letter': 'P',
        'icon': '🩹',
        'title': 'PAIN POINT',
        'subtitle': 'Differentiation',
        'question': 'Vitamin or Painkiller?',
        'leftLabel': 'Nice-to-have (UI)',
        'rightLabel': '10x Solution',
        'rubric': {
            'low': 'Cosmetic changes. Just a prettier app for a standard bank account.',
            'medium': 'Reduces friction (faster onboarding), but core product is standard.',
            'high': 'Solves deep friction (e.g., instant cross-border). Users cannot go back.'
        },
        'challengePrompts': ['Is it 10% better or 10x better?', 'What is the incumbent solution?']
    },
    {
        'id': 'automation',
        'letter': 'A',
        'icon': '🤖',
        'title': 'AUTOMATION',
        'subtitle': 'Tech Depth',
        'question': 'Wrapper or Deep Tech?',
        'leftLabel': 'Human/Manual',
        'rightLabel': 'AI/Algorithmic',
        'rubric': {
            'low': 'Manual processes behind scenes. Rule-based logic only.',
            'medium': 'Some automation in KYC, but support is human-heavy.',
            'high': 'Proprietary AI/ML models. Algorithmic underwriting. Self-driving finance.'
        },
        'challengePrompts': ['Real ML or just rules?', 'Proprietary tech ownership?']
    },
    {
        'id': 'compliance',
        'letter': 'C',
        'icon': '⚖️',
        'title': 'COMPLIANCE',
        'subtitle': 'Trust',
        'question': 'Regulatory Arbitrage or Trust?',
        'leftLabel': 'Grey Area',
        'rightLabel': 'Fully Licensed',
        'rubric': {
            'low': 'Unregulated. Operating across borders to avoid rules.',
            'medium': 'Partnering with a sponsor bank (BaaS) to rent a license.',
            'high': 'Full Banking Charter. Direct regulator relationship. Heavy compliance.'
        },
        'challengePrompts': ['Licenses held?', 'Regulatory scrutiny history?']
    },
    {
        'id': 'target',
        'letter': 'T',
        'icon': '🎯',
        'title': 'TARGET',
        'subtitle': 'Inclusion',
        'question': 'Mass Market or Niche?',
        'leftLabel': 'Mass Market',
        'rightLabel': 'Underserved Niche',
        'rubric': {
            'low': 'Competing for prime customers (High FICO) like major banks.',
            'medium': 'Millennials/Gen-Z focus, but still generally bankable.',
            'high': 'Unbanked, gig-workers, immigrants, or specific vertical niches.'
        },
        'challengePrompts': ['Who is excluded by banks?', 'Is the niche defensible?']
    }
]

# --- Helper Functions ---
def get_score_color(score):
    if score < 30: return "#ef4444"
    if score < 70: return "#f59e0b"
    return "#22c55e"

def get_score_label(score):
    if score < 30: return "Low Impact"
    if score < 70: return "Medium Impact"
    return "High Impact"

def create_radar_chart(values, show_benchmark):
    categories = [d['title'] for d in DIMENSIONS]
    
    r_values = list(values) + [values[0]]
    theta_values = categories + [categories[0]]
    
    fig = go.Figure()

    fig.add_trace(go.Scatterpolar(
        r=r_values,
        theta=theta_values,
        fill='toself',
        name='Your Analysis',
        line=dict(color='#667eea', width=3),
        fillcolor='rgba(102, 126, 234, 0.3)',
        marker=dict(size=8, color='#667eea')
    ))

    if show_benchmark:
        bank_values = [20, 80, 20, 30, 95, 20] 
        bank_r = bank_values + [bank_values[0]]
        
        fig.add_trace(go.Scatterpolar(
            r=bank_r,
            theta=theta_values,
            fill='toself',
            name='Traditional Bank',
            line=dict(color='#94a3b8', width=2, dash='dot'),
            fillcolor='rgba(148, 163, 184, 0.15)',
            marker=dict(size=6, color='#94a3b8')
        ))

    fig.update_layout(
        polar=dict(
            radialaxis=dict(
                visible=True,
                range=[0, 100],
                tickfont=dict(size=10),
                gridcolor='#e2e8f0'
            ),
            angularaxis=dict(
                gridcolor='#e2e8f0'
            )
        ),
        showlegend=True,
        legend=dict(
            orientation="h",
            yanchor="bottom",
            y=-0.15,
            xanchor="center",
            x=0.5,
            font=dict(size=11)
        ),
        margin=dict(l=60, r=60, t=40, b=60),
        height=450,
        paper_bgcolor="rgba(255,255,255,0.9)",
        plot_bgcolor="rgba(0,0,0,0)",
        font=dict(family="Inter, sans-serif")
    )
    return fig

def generate_word_doc(company_name, avg_score, timestamp):
    doc = Document()
    
    # Title
    heading = doc.add_heading('FINTECH IMPACT RADAR ANALYSIS', 0)
    heading.alignment = 1  # Center
    
    # Company info
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run(f"Company: {company_name if company_name else 'Not specified'}")
    run.bold = True
    run.font.size = Pt(14)
    
    p2 = doc.add_paragraph()
    run2 = p2.add_run(f"Analysis Date: {timestamp}")
    run2.font.size = Pt(10)
    run2.font.color.rgb = RGBColor(100, 116, 139)
    
    # Overall score
    doc.add_heading(f'Overall IMPACT Score: {avg_score}/100', level=1)
    score_label = get_score_label(avg_score)
    p3 = doc.add_paragraph(f"Assessment: {score_label}")
    p3.runs[0].italic = True
    
    doc.add_page_break()
    
    # Dimension details
    for dim in DIMENSIONS:
        score = st.session_state[f"score_{dim['id']}"]
        notes = st.session_state[f"note_{dim['id']}"]
        
        doc.add_heading(f"{dim['icon']} {dim['title']} - {score}/100", level=2)
        doc.add_paragraph(f"Focus: {dim['subtitle']}")
        doc.add_paragraph(f"Key Question: {dim['question']}")
        
        doc.add_heading('Scoring Rubric:', level=3)
        doc.add_paragraph(f"Low (0-30): {dim['rubric']['low']}")
        doc.add_paragraph(f"Medium (31-70): {dim['rubric']['medium']}")
        doc.add_paragraph(f"High (71-100): {dim['rubric']['high']}")
        
        doc.add_heading('Analysis & Evidence:', level=3)
        if notes:
            doc.add_paragraph(notes)
        else:
            doc.add_paragraph("No notes recorded.", style='Intense Quote')
        
        doc.add_paragraph("_" * 70)
        doc.add_paragraph()
    
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def export_json():
    """Export analysis as JSON"""
    data = {
        'company_name': st.session_state.company_name,
        'timestamp': datetime.now().isoformat(),
        'overall_score': round(sum([st.session_state[f"score_{d['id']}"] for d in DIMENSIONS]) / 6),
        'dimensions': {}
    }
    
    for dim in DIMENSIONS:
        data['dimensions'][dim['id']] = {
            'title': dim['title'],
            'score': st.session_state[f"score_{dim['id']}"],
            'notes': st.session_state[f"note_{dim['id']}"]
        }
    
    return json.dumps(data, indent=2)

def reset_state():
    st.session_state.company_name = ""
    for dim in DIMENSIONS:
        st.session_state[f"score_{dim['id']}"] = 50
        st.session_state[f"note_{dim['id']}"] = ""

# --- Initialize Session State ---
if 'company_name' not in st.session_state:
    st.session_state.company_name = ""
for dim in DIMENSIONS:
    if f"score_{dim['id']}" not in st.session_state:
        st.session_state[f"score_{dim['id']}"] = 50
    if f"note_{dim['id']}" not in st.session_state:
        st.session_state[f"note_{dim['id']}"] = ""

# --- SIDEBAR ---
with st.sidebar:
    st.title("📡 Analysis Controls")
    
    st.markdown("### 🏢 Company Information")
    st.text_input("Company Name", placeholder="e.g., Revolut, Stripe, Chime", key="company_name")
    
    # Calculate current metrics
    current_scores = [st.session_state[f"score_{d['id']}"] for d in DIMENSIONS]
    avg_score = round(sum(current_scores) / 6)
    
    # Quick stats
    st.markdown("### 📊 Quick Stats")
    col1, col2 = st.columns(2)
    with col1:
        st.metric("Avg Score", f"{avg_score}/100", delta=None)
    with col2:
        high_count = sum(1 for s in current_scores if s >= 70)
        st.metric("High Scores", f"{high_count}/6", delta=None)
    
    st.markdown("---")
    
    st.markdown("### 🔍 View Options")
    show_benchmark = st.checkbox("📈 Show Traditional Bank Benchmark", value=True)
    show_challenge_prompts = st.checkbox("💡 Show Challenge Prompts", value=False)
    
    st.markdown("---")
    
    st.markdown("### 📤 Export Analysis")
    
    timestamp = datetime.now().strftime("%Y-%m-%d %H:%M")
    
    # Word document
    docx_file = generate_word_doc(st.session_state.company_name, avg_score, timestamp)
    st.download_button(
        label="📄 Download Report (DOCX)", 
        data=docx_file,
        file_name=f"IMPACT_Analysis_{st.session_state.company_name or 'Company'}_{datetime.now().strftime('%Y%m%d')}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        use_container_width=True
    )
    
    # JSON export
    json_data = export_json()
    st.download_button(
        label="📋 Download Data (JSON)",
        data=json_data,
        file_name=f"IMPACT_Data_{st.session_state.company_name or 'Company'}_{datetime.now().strftime('%Y%m%d')}.json",
        mime="application/json",
        use_container_width=True
    )
    
    st.markdown("---")
    
    if st.button("🔄 Reset All Scores", use_container_width=True, type="secondary"):
        reset_state()
        st.rerun()

    st.markdown("---")
    
    with st.expander("📚 IMPACT Framework Guide"):
        for dim in DIMENSIONS:
            st.markdown(f"**{dim['icon']} {dim['title']}**")
            st.caption(f"{dim['subtitle']} - {dim['question']}")
            st.markdown("")

# --- MAIN LAYOUT ---

# Header
st.title("The Fintech IMPACT Radar")
st.markdown(f"<p style='font-size: 1.1rem; color: white; margin-top: -0.5rem;'>Comprehensive analysis framework for evaluating fintech innovation and disruption potential</p>", unsafe_allow_html=True)

st.markdown("<br>", unsafe_allow_html=True)

# Top Section: Chart and Summary
col_chart, col_summary = st.columns([1.4, 1.6])

with col_chart:
    st.markdown("### 📊 Visual Profile")
    radar_fig = create_radar_chart(current_scores, show_benchmark)
    st.plotly_chart(radar_fig, use_container_width=True)

with col_summary:
    st.markdown("### 🎯 Impact Summary")
    
    # Overall score card
    score_color = get_score_color(avg_score)
    st.markdown(f"""
        <div style='background: linear-gradient(135deg, {score_color}15 0%, {score_color}30 100%); 
                    padding: 1.5rem; border-radius: 12px; border-left: 4px solid {score_color};
                    margin-bottom: 1rem;'>
            <div style='font-size: 0.9rem; color: #64748b; font-weight: 600;'>OVERALL IMPACT SCORE</div>
            <div style='font-size: 3rem; font-weight: 700; color: {score_color}; line-height: 1.2;'>{avg_score}<span style='font-size: 1.5rem;'>/100</span></div>
            <div style='font-size: 0.95rem; color: #334155; font-weight: 600; margin-top: 0.5rem;'>{get_score_label(avg_score)}</div>
        </div>
    """, unsafe_allow_html=True)
    
    # Dimension breakdown
    st.markdown("**Dimension Scores:**")
    for dim in DIMENSIONS:
        score = st.session_state[f"score_{dim['id']}"]
        color = get_score_color(score)
        st.markdown(f"""
            <div style='display: flex; justify-content: space-between; align-items: center; 
                        padding: 0.75rem; background: white; border-radius: 8px; 
                        margin-bottom: 0.5rem; box-shadow: 0 1px 3px rgba(0,0,0,0.08);'>
                <span style='font-weight: 600; color: #334155;'>{dim['icon']} {dim['title']}</span>
                <span style='background: {color}; color: white; padding: 4px 12px; 
                             border-radius: 12px; font-weight: 700; font-size: 0.9rem;'>{score}</span>
            </div>
        """, unsafe_allow_html=True)

st.markdown("<br>", unsafe_allow_html=True)

# Dimension Analysis Grid
st.markdown("### 🔍 Detailed Analysis Grid")
st.markdown("<p style='color: white; margin-top: -0.5rem;'>Score each dimension and document your evidence below</p>", unsafe_allow_html=True)

# Create grid
for i in range(0, 6, 2):
    cols = st.columns(2)
    for j, col in enumerate(cols):
        if i + j < 6:
            dim = DIMENSIONS[i + j]
            with col:
                with st.container(border=True):
                    # Header
                    c_head, c_val = st.columns([3.5, 1])
                    with c_head:
                        st.markdown(f"### {dim['icon']} {dim['title']}")
                        st.caption(dim['subtitle'])
                    with c_val:
                        score = st.session_state[f"score_{dim['id']}"]
                        color = get_score_color(score)
                        st.markdown(f"<div class='score-badge' style='background: {color}; color: white;'>{score}</div>", unsafe_allow_html=True)
                    
                    # Question
                    st.markdown(f"<div style='background: #f8fafc; padding: 0.75rem; border-radius: 8px; margin: 0.75rem 0; border-left: 3px solid #667eea;'><b>❓ {dim['question']}</b></div>", unsafe_allow_html=True)
                    
                    # Slider
                    st.slider(
                        "Score", 0, 100, 
                        key=f"score_{dim['id']}", 
                        label_visibility="collapsed"
                    )
                    
                    # Labels
                    l, r = st.columns(2)
                    l.caption(f"◀ {dim['leftLabel']}")
                    r.markdown(f"<div style='text-align:right; font-size: 0.8rem; color: #64748b; font-weight: 500;'>{dim['rightLabel']} ▶</div>", unsafe_allow_html=True)
                    
                    # Challenge prompts
                    if show_challenge_prompts:
                        st.markdown("<div style='margin-top: 0.75rem;'>", unsafe_allow_html=True)
                        for prompt in dim['challengePrompts']:
                            st.markdown(f"<span class='info-badge'>💭 {prompt}</span>", unsafe_allow_html=True)
                        st.markdown("</div>", unsafe_allow_html=True)
                    
                    # Expander
                    with st.expander("📋 View Rubric & Add Notes"):
                        st.markdown("#### 📏 Scoring Rubric")
                        st.markdown(f"""
                        <table class="rubric-table">
                            <tr><td class="rubric-header">Low<br>(0-30)</td><td>{dim['rubric']['low']}</td></tr>
                            <tr><td class="rubric-header">Medium<br>(31-70)</td><td>{dim['rubric']['medium']}</td></tr>
                            <tr><td class="rubric-header">High<br>(71-100)</td><td>{dim['rubric']['high']}</td></tr>
                        </table>
                        """, unsafe_allow_html=True)
                        
                        st.markdown("#### 📝 Evidence & Notes")
                        st.text_area(
                            "Document your reasoning", 
                            key=f"note_{dim['id']}", 
                            height=100, 
                            placeholder="Why did you assign this score? What evidence supports it?",
                            label_visibility="collapsed"
                        )

# Footer
st.markdown("<br><br>", unsafe_allow_html=True)
st.markdown("---")
st.markdown("""
    <div style='text-align: center; padding: 1rem; color: white;'>
        <p style='font-size: 0.9rem; margin: 0;'>Built by Professor Vangelis Tsiligkiris | 
        <a href='#' style='color: white; text-decoration: underline;'>Learn More</a> about the IMPACT Framework</p>
    </div>
""", unsafe_allow_html=True)